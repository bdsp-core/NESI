# build_table_1_docx.py
"""
Build a Word document containing Table 1 for the encephalopathy index paper.

Reads:
  - exposure_summary_by_patient_final.csv
  - exposure_summary_by_eeg_final.csv

Produces a single .docx file with:
  - Table caption
  - Two-section Table 1: PATIENT-LEVEL EXPOSURE and EEG-LEVEL EXPOSURE
  - Explicit section header rows for "Sedative infusion" and "Opiate infusion"
    (header rows have no data, just visual grouping)
  - Indented sub-rows for top-3 ASMs, sedative drugs, opiate infusion drugs
  - Footnotes for parenteral/enteral benzodiazepine, parenteral/enteral opiate,
    and propofol (bolus + infusion lumped)

Output: I0001_ASM_ANTIPSYCH_DIR/Table_1_medication_exposure.docx
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

import med_config

try:
    import winsound
    HAS_WINSOUND = True
except ImportError:
    HAS_WINSOUND = False


# ─── Display labels for each Category, plus footnote markers ──────────────────
# Maps the CSV "Category" value to (display_text, footnote_marker_or_None).
# If display_text is None, use the CSV value as-is.

DISPLAY_RENAMES = {
    "Parenteral benzo (non-infusion)":  ("Parenteral benzodiazepine",   "a"),
    "Enteral benzo (non-ASM)":          ("Enteral benzodiazepine",      "b"),
    "Parenteral opiate":                ("Parenteral opiate",           "c"),
    "Enteral opiate":                   ("Enteral opiate",              "d"),
    "Antipsychotic":                    ("Antipsychotics (any route)",  None),
    "  Propofol":                       ("  Propofol",                  "e"),
}

# Categories that are sub-rows (indented)
SUB_ROWS = {
    "  Levetiracetam", "  Lacosamide", "  Valproate",
    "  Propofol", "  Midazolam", "  Ketamine", "  Dexmedetomidine",
    "  Fentanyl", "  Morphine", "  Hydromorphone",
}

# Rows that are bolded section headers WITH data (e.g. "Maintenance ASM")
TOP_LEVEL_ROWS = {
    "Maintenance ASM",
    "Parenteral benzo (non-infusion)",
    "Enteral benzo (non-ASM)",
    "Sedative infusion",
    "Opiate infusion",
    "Parenteral opiate",
    "Enteral opiate",
    "Antipsychotic",
}


def add_styled_run(paragraph, text, bold=False, size=10, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return run


def style_cell(cell, text, bold=False, indent=False, superscript_suffix=None,
               align="left", italic=False):
    cell.text = ""
    para = cell.paragraphs[0]
    if align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    display = text.lstrip() if indent else text
    if indent:
        para.paragraph_format.left_indent = Inches(0.2)

    add_styled_run(para, display, bold=bold, italic=italic)

    if superscript_suffix:
        run = para.add_run(superscript_suffix)
        run.font.superscript = True
        run.font.size = Pt(10)

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_section(doc, df, section_title):
    """
    Build one section (patient-level or EEG-level) as a Word table.
    df has columns: Category + cohort labels. First row is the 'Total ___' row.
    """
    h = doc.add_paragraph()
    add_styled_run(h, section_title, bold=True, size=11)

    cohort_cols = [c for c in df.columns if c != "Category"]

    # Build the row list: insert section headers ("Sedative infusion", "Opiate
    # infusion") before their anchor rows so they appear as separate visual rows.
    rows_to_render = [{"_data": row} for _, row in df.iterrows()]
    n_data_rows = len(rows_to_render)
    n_cols = 1 + len(cohort_cols)
    n_rows = 1 + n_data_rows

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # Header row (cohort column labels)
    hdr = table.rows[0].cells
    style_cell(hdr[0], "", bold=True)
    for i, c in enumerate(cohort_cols):
        style_cell(hdr[i + 1], c, bold=True, align="center")

    # Data + section-header rows
    for r_idx, item in enumerate(rows_to_render):
        cells = table.rows[r_idx + 1].cells



        row = item["_data"]
        csv_cat = str(row["Category"])

        # Resolve display label + footnote marker
        if csv_cat in DISPLAY_RENAMES:
            display_label, sup = DISPLAY_RENAMES[csv_cat]
        else:
            display_label, sup = csv_cat, None

        is_sub = csv_cat in SUB_ROWS
        is_top_level = csv_cat in TOP_LEVEL_ROWS
        is_total = csv_cat.startswith("Total")
        bold = is_total or is_top_level

        style_cell(
            cells[0], display_label,
            bold=bold,
            indent=is_sub,
            superscript_suffix=sup,
            align="left",
        )
        for i, c in enumerate(cohort_cols):
            style_cell(cells[i + 1], str(row[c]), align="center")


try:
    pt_path  = med_config.I0001_ASM_ANTIPSYCH_DIR / "exposure_summary_by_patient_final.csv"
    eeg_path = med_config.I0001_ASM_ANTIPSYCH_DIR / "exposure_summary_by_eeg_final.csv"

    if not pt_path.exists() or not eeg_path.exists():
        raise FileNotFoundError(
            f"Required CSVs missing. Run build_final_exposure_tables.py first.\n"
            f"  Expected: {pt_path}\n"
            f"  Expected: {eeg_path}"
        )

    pt_df  = pd.read_csv(pt_path)
    eeg_df = pd.read_csv(eeg_path)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title / caption
    title = doc.add_paragraph()
    add_styled_run(title, "Table 1. ", bold=True, size=11)
    add_styled_run(
        title,
        "Medication exposure across cohorts. "
        "Exposure was defined as any administration of the indicated drug "
        "or drug class within the 24 hours preceding the relevant EEG/score "
        "timestamp. The top section reports the proportion of unique patients "
        "with any qualifying exposure; the bottom section reports the "
        "proportion of EEG/score timepoints with qualifying exposure.",
        size=11,
    )

    # Patient-level section
    build_section(doc, pt_df, "Patient-level exposure")
    doc.add_paragraph()

    # EEG-level section
    build_section(doc, eeg_df, "EEG-level exposure")

    # ─── Footnotes ────────────────────────────────────────────────────────────
    footnotes = [
        ("a",
         "Parenteral benzodiazepine: bolus administration of lorazepam, "
         "diazepam, midazolam (when not given as a continuous infusion), or "
         "other benzodiazepines via intravenous, intramuscular, or "
         "intranasal routes. These routes were combined because they all "
         "produce clinically equivalent CNS exposure within the 24-hour "
         "exposure window."),
        ("b",
         "Enteral benzodiazepine: lorazepam, diazepam, midazolam, or other "
         "benzodiazepines administered via oral, sublingual, rectal, or "
         "buccal routes. Excludes long-acting maintenance ASMs (clonazepam, "
         "clobazam) which are counted under maintenance ASM."),
        ("c",
         "Parenteral opiate: bolus administration of fentanyl, morphine, "
         "or hydromorphone via intravenous, intramuscular, or intranasal "
         "routes."),
        ("d",
         "Enteral opiate: fentanyl, morphine, or hydromorphone "
         "administered via oral, sublingual, transdermal, rectal, or buccal "
         "routes."),
        ("e",
         "Propofol exposure includes both continuous infusion and bolus "
         "administration (e.g., procedural sedation) within the 24-hour "
         "exposure window."),
    ]
    for marker, text in footnotes:
        para = doc.add_paragraph()
        run = para.add_run(marker)
        run.font.superscript = True
        run.font.size = Pt(9)
        rest = para.add_run(" " + text)
        rest.font.size = Pt(9)

    out_path = med_config.I0001_ASM_ANTIPSYCH_DIR / "Table_1_medication_exposure.docx"
    doc.save(out_path)
    print(f"Wrote: {out_path}")

    print("\nDone.")
    if HAS_WINSOUND:
        winsound.MessageBeep(winsound.MB_OK)
    else:
        try:
            print("\a")
        except Exception:
            pass

except Exception as e:
    print(f"\nERROR: {e}")
    if HAS_WINSOUND:
        winsound.MessageBeep(winsound.MB_ICONHAND)
    else:
        try:
            print("\a")
        except Exception:
            pass
    raise